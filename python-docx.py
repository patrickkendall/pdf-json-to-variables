from docx import Document
from docx.shared import Inches
from PIL import Image
import requests
from docx.shared import Pt

document = Document()

font = document.styles['Normal'].font
font.name = 'Arial'
font.size = Pt(24)

p = document.add_paragraph('Title')

document.add_picture('bulba.png', width=Inches(3))
document.add_picture('bulba.png', width=Inches(3))
document.add_picture('bulba.png', width=Inches(3))

document.save('demo.docx')